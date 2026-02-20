import os
import re
import sys
import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)

# Configuration
ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUTPUT_FILE = os.path.join(ROOT_DIR, "auditoria_codigo_completa.docx")

EXCLUDE_DIRS = {
    "node_modules",
    ".git",
    "__pycache__",
    "dist",
    "build",
    ".vscode",
    ".idea",
    "venv",
    "env",
    "tmp",
    "coverage",
    ".next",
    # Removed: "motores_internos", "site", "conocimiento", "tools" to include EVERYTHING
}

EXCLUDE_FILES = {
    "package-lock.json",
    "yarn.lock",
    ".DS_Store",
    "auditoria_completa_con_explicaciones.md",
    "tree_output.txt",
    "stats.json",
    "generar_auditoria_word.py",
    "txt_to_docx.py",
    "auditoria_codigo_completa.docx",
    "estructura_completa.docx",
}

EXCLUDE_EXTENSIONS = {
    ".pyc",
    ".png",
    ".jpg",
    ".jpeg",
    ".gif",
    ".ico",
    ".svg",
    ".woff",
    ".woff2",
    ".ttf",
    ".eot",
    ".mp4",
    ".mov",
    ".zip",
    ".tar",
    ".gz",
    ".rar",
    ".7z",
    ".pdf",
    ".exe",
    ".dll",
    ".bin",
    ".dat",
    ".db",
    ".sqlite",
    ".sqlite3",
    ".pkl",
}

MAX_FILE_SIZE_BYTES = 1 * 1024 * 1024  # 1MB limit for code files


def is_text_file(filename):
    return not any(filename.lower().endswith(ext) for ext in EXCLUDE_EXTENSIONS)


def get_child_explanation(filepath, content):
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filename)[1].lower()

    # "Explanation for a 3-year-old" - First Person
    explanation = "Aquí he guardado un archivo del sistema para que todo funcione."

    if ext == ".py":
        explanation = "En este archivo le doy órdenes a mi robot. Le digo paso a paso lo que tiene que hacer con mis juguetes (datos)."
        if "main" in filename or "principal" in filename:
            explanation = "¡Este es el archivo jefe! Desde aquí arranco todo mi juego."
        elif "config" in filename:
            explanation = "Aquí guardo mis secretos y las reglas del juego."
        elif "util" in filename or "tool" in filename:
            explanation = "Estas son mis herramientas magicas para arreglar cosas."

    elif ext in [".js", ".jsx", ".ts", ".tsx"]:
        explanation = "Aquí pinto cosas bonitas en la pantalla para que las veas."
        if "component" in filepath:
            explanation = (
                "He creado una pieza de Lego para construir mi casita en la web."
            )

    elif ext == ".css":
        explanation = "Aquí elijo los colores de mi dibujo. ¡Mira qué bonito queda!"

    elif ext == ".html":
        explanation = (
            "Este es el esqueleto de mi página web. Son los huesos que sujetan todo."
        )

    elif ext == ".json":
        explanation = (
            "Aquí guardo mis cromos y mis cosas importantes en cajitas ordenadas."
        )

    elif ext in [".bat", ".cmd", ".sh"]:
        explanation = (
            "¡Boton rojo! Con esto hago que el ordenador arranque a toda velocidad."
        )

    return explanation


def add_line_comments(content, ext):
    """
    Adds simple Spanish comments to each line based on rudimentary heuristics.
    """
    lines = content.splitlines()
    commented_lines = []

    # Simple dictionary of keywords -> comments
    # Tono: "Yo hago esto"
    keywords = {
        "import ": " # Traigo mi caja de herramientas",
        "from ": " # Saco una herramienta de la caja",
        "def ": " # Voy a enseñarle un truco nuevo a mi robot",
        "class ": " # Creo una fábrica de juguetes",
        "return ": " # Te devuelvo el resultado",
        "if ": " # Si pasa esto...",
        "else": " # Si no, hago esto otro",
        "elif": " # O quizás pase esto...",
        "for ": " # Repito esto muchas veces con mis juguetes",
        "while ": " # Sigo haciendo esto mientras no me digas que pare",
        "try:": " # Voy a intentar hacer esto con cuidado",
        "except": " # ¡Ups! Si algo sale mal, lo arreglo aquí",
        "print": " # Escribo esto en la pantalla",
        "console.log": " # Le digo hola a la consola",
        "const ": " # Esto no va a cambiar nunca",
        "let ": " # Esto puede cambiar luego",
        "var ": " # Una cajita para guardar cosas",
        "await ": " # Espero un poquito a que termine",
        "async ": " # Lo hago rápido sin esperar",
        "True": " # ¡Es verdad!",
        "False": " # ¡Es mentira!",
        "None": " # No hay nada aquí",
        "null": " # Está vacío",
        "self.": " # Esto es mío",
        "this.": " # Esto es mío",
        "pass": " # Aquí no hago nada",
        "break": " # ¡Paro ya!",
        "continue": " # Sigo con el siguiente",
        "open": " # Abro un libro para leer",
        "close": " # Cierro el libro",
        "write": " # Escribo en mi cuaderno",
        "read": " # Leo lo que pone",
        "path": " # Busco el camino",
        "os.": " # Le hablo al ordenador",
        "sys.": " # Toco los botones del sistema",
        "json.": " # Ordeno mis cromos",
    }

    for line in lines:
        stripped = line.strip()

        # Skip empty lines or existing comments (mostly)
        if not stripped:
            commented_lines.append(line)
            continue

        # Avoid double commenting if line already has a comment
        if "#" in line or "//" in line:
            commented_lines.append(line)
            continue

        # Heuristic matching
        added_comment = ""

        # Check for assignment specifically
        if "=" in stripped and not "==" in stripped and not added_comment:
            added_comment = " # Guardo esto en una cajita nueva"

        # Check keywords
        for k, v in keywords.items():
            if k in stripped:
                added_comment = v
                break  # Take the first match

        # Default fallback for lines with code match regex for function call
        if not added_comment and re.search(r"\w+\(", stripped):
            added_comment = " # Llamo a uno de mis amigos para que haga algo"

        # Apply comment logic based on extension logic
        if added_comment:
            # Python-style comments
            if ext in [".py", ".sh", ".yaml", ".yml"]:
                commented_lines.append(f"{line:<60}{added_comment}")
            # C-style comments (JS, C, etc)
            elif ext in [
                ".js",
                ".jsx",
                ".ts",
                ".tsx",
                ".css",
                ".java",
                ".c",
                ".cpp",
                ".h",
            ]:
                # Convert # to //
                c_comment = added_comment.replace("#", "//")
                commented_lines.append(f"{line:<60}{c_comment}")
            # Batch
            elif ext in [".bat", ".cmd"]:
                # Batch syntax is tricky for inline, usually :: or REM but inline & :: often fails using REM is safer on new line
                # For simplicity, we just won't touch batch inline heavily or use REM at end if supported (often not well)
                # Let's skip inline for BAT to avoid breaking it, just append line.
                commented_lines.append(line)
            # HTML/XML
            elif ext in [".html", ".xml"]:
                html_comment = (
                    added_comment.replace("#", "<!--").replace("\n", "") + " -->"
                )
                commented_lines.append(f"{line:<60}{html_comment}")
            else:
                commented_lines.append(line)
        else:
            commented_lines.append(line)

    return "\n".join(commented_lines)


def add_file_to_docx(doc, filepath, relpath):
    filename = os.path.basename(filepath)
    ext = os.path.splitext(filename)[1].lower()

    try:
        if os.path.getsize(filepath) > MAX_FILE_SIZE_BYTES:
            doc.add_heading(f"📄 {relpath} (Muy Grande)", level=2)
            doc.add_paragraph("Este libro es demasiado gordo. No me cabe aquí.")
            return

        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        explanation = get_child_explanation(filepath, content)
        commented_content = add_line_comments(content, ext)

        # Heading
        heading = doc.add_heading(f"📄 {relpath}", level=2)

        # Explanation
        p_expl = doc.add_paragraph()
        run_expl = p_expl.add_run(f"YO DIGO: {explanation}")
        run_expl.bold = True
        run_expl.font.color.rgb = RGBColor(0, 112, 192)  # Nice Blue

        # Code Content
        doc.add_paragraph("MI CÓDIGO:", style="Normal")

        # Add code in monospaced font, small size
        p_code = doc.add_paragraph(commented_content)
        p_code.style = doc.styles["Normal"]
        p_code.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p_code.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(8)

    except Exception as e:
        print(f"Error processing {relpath}: {e}")
        doc.add_paragraph(f"¡Ups! No he podido leer esto: {e}")


def main():
    print("Iniciando generación de auditoría Word 'Modo Niño'...")

    doc = Document()

    # Title
    title = doc.add_heading("MIS COSAS DE ZEROX (EXPLICADO PARA TI)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Hola. Aquí te enseño todo lo que he construido. Te lo explico fácil para que lo entiendas."
    )

    # --- INSERTAR ÍNDICE (TREE) ---
    tree_file = os.path.join(ROOT_DIR, "tree_output.txt")
    if os.path.exists(tree_file):
        print(f"Insertando índice desde {tree_file}...")
        doc.add_page_break()
        doc.add_heading("ASÍ ESTÁ ORDENADO TODO (INDICE)", level=1)

        try:
            with open(tree_file, "r", encoding="utf-8", errors="replace") as f:
                tree_content = f.read()

            p_tree = doc.add_paragraph(tree_content)
            p_tree.style = doc.styles["Normal"]
            for run in p_tree.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8)

            print("Índice insertado.")

        except Exception as e:
            print(f"Error insertando índice: {e}")
            doc.add_paragraph(f"No pude poner el índice: {e}")
    else:
        print("⚠️ No encontré el archivo del índice (tree_output.txt).")

    doc.add_page_break()
    # ------------------------------

    file_count = 0

    for root, dirs, files in os.walk(ROOT_DIR):
        dirs.sort()
        files.sort()
        dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]

        for file in files:
            if file in EXCLUDE_FILES or not is_text_file(file):
                continue

            filepath = os.path.join(root, file)
            relpath = os.path.relpath(filepath, ROOT_DIR)

            print(f"Procesando: {relpath}")
            add_file_to_docx(doc, filepath, relpath)
            file_count += 1

    print(f"Guardando documento con {file_count} archivos...")
    try:
        doc.save(OUTPUT_FILE)
        print(f"✅ Guardado en: {OUTPUT_FILE}")
    except PermissionError:
        print(f"⚠️ El archivo {OUTPUT_FILE} está ocupado.")
        new_output = OUTPUT_FILE.replace(
            ".docx", f"_{datetime.datetime.now().strftime('%H%M%S')}.docx"
        )
        doc.save(new_output)
        print(f"✅ Guardado en: {new_output}")


if __name__ == "__main__":
    main()
