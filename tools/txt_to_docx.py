import os
import sys

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Please install python-docx: pip install python-docx")
    sys.exit(1)


def main():
    input_file = "tree_output.txt"
    # Ensure output directory exists
    output_dir = (
        "c:/Users/manue/.gemini/antigravity/brain/7a551a2b-7d1b-4cb4-9b24-d4adaf0d4fdb"
    )
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    output_file = os.path.join(output_dir, "estructura_completa.docx")

    if not os.path.exists(input_file):
        print(f"Error: {input_file} not found.")
        return

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Consolas"
    font.size = Pt(9)
    # Just in case, try to set paragraph format to single spacing
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(0)

    doc.add_heading("Estructura Completa del Proyecto ZEROX", 0)

    print(f"Reading {input_file}...")
    try:
        # Use iterator to read file line by line to save memory
        with open(input_file, "r", encoding="utf-8", errors="replace") as f:
            # Setup a paragraph for accumulation
            current_buffer = []
            buffer_size = 500  # lines per paragraph block to speed up

            count = 0
            for line in f:
                current_buffer.append(line)
                if len(current_buffer) >= buffer_size:
                    # Join and add
                    text_block = "".join(current_buffer)
                    p = doc.add_paragraph(text_block)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_buffer = []
                    count += 1
                    if count % 10 == 0:
                        print(".", end="", flush=True)  # Progress indicator

            # Add remaining
            if current_buffer:
                text_block = "".join(current_buffer)
                p = doc.add_paragraph(text_block)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        print("\nSaving document...")
        doc.save(output_file)
        print(f"Saved to {output_file}")

    except Exception as e:
        print(f"Error: {e}")


if __name__ == "__main__":
    main()
